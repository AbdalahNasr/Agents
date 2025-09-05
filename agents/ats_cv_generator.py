#!/usr/bin/env python3
"""
ATS CV Generator - Creates ATS-optimized CVs based on job requirements
Integrates with the ATS optimizer to generate job-specific CVs

Features:
- Job-specific CV customization
- ATS-friendly formatting
- Keyword optimization
- Contact information validation
- Multiple output formats (PDF, DOCX, TXT, HTML)
- Real-time ATS scoring
"""

import os
import json
from datetime import datetime
from typing import Dict, List, Any, Optional
from pathlib import Path

import sys
sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from utils.logger import AgentLogger
from agents.ats_optimizer import ATSOptimizer
from config import Config

class ATSCVGenerator:
    """ATS-optimized CV generator"""
    
    def __init__(self):
        """Initialize the ATS CV Generator"""
        self.logger = AgentLogger("ats_cv_generator")
        self.ats_optimizer = ATSOptimizer()
        
        # ATS-friendly CV template
        self.cv_template = {
            "header": {
                "name": "ABDALLAH NASR ALI",
                "title": "JUNIOR FULL STACK DEVELOPER",
                "email": "body16nasr16bn@gmail.com",
                "phone": "+20 106 950 9757",
                "address": "Cairo, Egypt",
                "linkedin": "https://linkedin.com/in/abdallah-nasr-ali",
                "github": "https://github.com/AbdalahNasr",
                "portfolio": "https://my-v3-potfolio.vercel.app"
            },
            "summary": "Passionate Junior Full Stack Developer seeking first professional opportunity. Strong foundation in modern web technologies with hands-on experience building responsive web applications, RESTful APIs, and user-friendly interfaces. Eager to contribute to a development team and grow professionally. Committed to writing clean, maintainable code and continuously learning new technologies.",
            "skills": {
                "frontend": ["React", "Angular", "JavaScript", "TypeScript", "HTML5", "CSS3", "SCSS", "Tailwind CSS", "Bootstrap", "Next.js", "jQuery"],
                "backend": ["Node.js", "Express.js", "REST APIs"],
                "database": ["MongoDB", "Prisma"],
                "tools": ["Git", "GitHub", "Vercel", "Socket.IO", "Redux", "NextAuth"],
                "other": ["JWT", "Swagger", "Cloudinary", "i18n", "PostCSS", "ESLint", "SSR"]
            },
            "experience": [
                {
                    "title": "Angular Developer Intern",
                    "company": "Link Data Center",
                    "duration": "08/2024 - 09/2024",
                    "description": "Integrated APIs, built reusable UI components, and created dynamic layouts with Angular. Used Git and version control in a team environment.",
                    "achievements": [
                        "Built responsive Angular e-commerce app with product filtering and cart functionality",
                        "Integrated REST APIs and deployed applications via Vercel",
                        "Gained hands-on experience with version control and team collaboration"
                    ]
                },
                {
                    "title": "Front End Angular Developer",
                    "company": "Angular E-Commerce App",
                    "duration": "11/2024",
                    "description": "Built a responsive Angular e-commerce app with product filtering, cart, checkout, and authentication. Integrated REST APIs and deployed via Vercel.",
                    "achievements": [
                        "Developed complete e-commerce application with authentication system",
                        "Implemented product filtering, cart, and checkout functionality",
                        "Successfully deployed application using Vercel platform"
                    ]
                }
            ],
            "education": [
                {
                    "degree": "Bachelor's in Management Information Systems",
                    "institution": "Cairo Higher Institute",
                    "duration": "07/2023",
                    "description": "Management Information Systems"
                },
                {
                    "degree": "Full Stack Web Development Diploma",
                    "institution": "Route Academy, Cairo",
                    "duration": "10/2023",
                    "description": "Full Stack Web Development"
                }
            ],
            "projects": [
                {
                    "name": "Threads App - Social Platform",
                    "description": "Next.js, TypeScript, Tailwind, Node.js. Auth system, post creation, likes, responsive design. Complete with working code and live demo.",
                    "technologies": ["Next.js", "TypeScript", "Tailwind", "Node.js"],
                    "achievements": ["Built complete social platform with authentication", "Implemented post creation and likes functionality", "Deployed live demo with responsive design"],
                    "code_url": "https://github.com/AbdalahNasr/threads",
                    "live_url": "https://threads-4cls.vercel.app/sign-in"
                },
                {
                    "name": "Order-Food App - Full-Stack Platform",
                    "description": "Next.js, TypeScript, Tailwind, Redux, Prisma, NextAuth. Admin dashboard, i18n, Cloudinary, product/cart management. Complete with working code.",
                    "technologies": ["Next.js", "TypeScript", "Tailwind", "Redux", "Prisma", "NextAuth"],
                    "achievements": ["Built complete food ordering platform with admin dashboard", "Implemented internationalization and image management", "Created product and cart management system"],
                    "code_url": "https://github.com/AbdalahNasr/order-food-app",
                    "live_url": None
                },
                {
                    "name": "ISTQP Quiz App - Quiz Platform",
                    "description": "Next.js, TypeScript, Tailwind CSS, PostCSS. JSON-based quizzes, multilingual support, ESLint, SSR. Complete with working code and live demo.",
                    "technologies": ["Next.js", "TypeScript", "Tailwind CSS", "PostCSS"],
                    "achievements": ["Developed multilingual quiz system with working code", "Implemented JSON-based quiz system with SSR", "Built responsive design with modern UI/UX"],
                    "code_url": "https://github.com/AbdalahNasr/istqp-quiz",
                    "live_url": "https://istqp-quiz.vercel.app"
                }
            ],
            "certificates": [
                {
                    "name": "Foundations of UX Design",
                    "issuer": "Google",
                    "date": "2024"
                },
                {
                    "name": "React Basics",
                    "issuer": "Google",
                    "date": "2024"
                },
                {
                    "name": "Full Stack Web Development",
                    "issuer": "Route Academy",
                    "date": "2023"
                },
                {
                    "name": "Angular Developer Internship",
                    "issuer": "Link Data Center",
                    "date": "2024"
                }
            ],
            "languages": [
                {
                    "language": "English",
                    "level": "Upper Intermediate"
                },
                {
                    "language": "Arabic",
                    "level": "Native"
                }
            ]
        }
        
        self.logger.info("ATS CV Generator initialized successfully")
    
    def generate_job_specific_cv(self, job_info: Dict[str, Any]) -> Dict[str, Any]:
        """
        Generate a job-specific ATS-optimized CV
        
        Args:
            job_info: Job information including title, company, description
            
        Returns:
            Dictionary containing optimized CV data and ATS analysis
        """
        self.logger.info(f"Generating job-specific CV for: {job_info.get('title', 'Unknown')} at {job_info.get('company', 'Unknown')}")
        
        try:
            # Start with base template
            cv_data = self.cv_template.copy()
            
            # Optimize for specific job
            cv_data = self._optimize_for_job(cv_data, job_info)
            
            # Generate CV content
            cv_content = self._generate_cv_content(cv_data)
            
            # Perform ATS analysis
            ats_result = self.ats_optimizer.analyze_cv_for_ats(
                cv_content=cv_content,
                job_description=job_info.get('description', ''),
                job_title=job_info.get('title', ''),
                company=job_info.get('company', '')
            )
            
            # Create result
            result = {
                "job_info": job_info,
                "cv_data": cv_data,
                "cv_content": cv_content,
                "ats_analysis": {
                    "overall_score": ats_result.overall_score,
                    "contact_info_score": ats_result.contact_info_score,
                    "job_title_match_score": ats_result.job_title_match_score,
                    "skill_match_score": ats_result.skill_match_score,
                    "formatting_score": ats_result.formatting_score,
                    "readability_score": ats_result.readability_score,
                    "web_presence_score": ats_result.web_presence_score,
                    "recommendations": ats_result.recommendations,
                    "missing_elements": ats_result.missing_elements,
                    "keyword_suggestions": ats_result.keyword_suggestions
                },
                "generated_at": datetime.now().isoformat()
            }
            
            self.logger.info(f"Job-specific CV generated successfully. ATS Score: {ats_result.overall_score}/100")
            return result
            
        except Exception as e:
            self.logger.error(f"Error generating job-specific CV: {e}")
            return None
    
    def _optimize_for_job(self, cv_data: Dict[str, Any], job_info: Dict[str, Any]) -> Dict[str, Any]:
        """Optimize CV data for specific job requirements"""
        job_title = job_info.get('title', '').lower()
        job_description = job_info.get('description', '').lower()
        
        # Update title to match job (but keep it realistic for entry-level)
        if 'senior' in job_title:
            cv_data['header']['title'] = 'JUNIOR FULL STACK DEVELOPER'  # Be honest about level
        elif 'lead' in job_title:
            cv_data['header']['title'] = 'JUNIOR FULL STACK DEVELOPER'  # Be honest about level
        elif 'frontend' in job_title:
            cv_data['header']['title'] = 'JUNIOR FRONTEND DEVELOPER'
        elif 'backend' in job_title:
            cv_data['header']['title'] = 'JUNIOR BACKEND DEVELOPER'
        elif 'mobile' in job_title:
            cv_data['header']['title'] = 'JUNIOR MOBILE DEVELOPER'
        elif 'junior' in job_title or 'entry' in job_title or 'trainee' in job_title:
            cv_data['header']['title'] = 'JUNIOR FULL STACK DEVELOPER'
        
        # Optimize summary for job
        cv_data['summary'] = self._optimize_summary(cv_data['summary'], job_info)
        
        # Add job-specific keywords to skills
        cv_data['skills'] = self._optimize_skills(cv_data['skills'], job_description)
        
        # Optimize experience descriptions
        cv_data['experience'] = self._optimize_experience(cv_data['experience'], job_description)
        
        # Optimize projects
        cv_data['projects'] = self._optimize_projects(cv_data['projects'], job_description)
        
        return cv_data
    
    def _optimize_summary(self, current_summary: str, job_info: Dict[str, Any]) -> str:
        """Optimize summary for specific job"""
        job_title = job_info.get('title', '')
        job_description = job_info.get('description', '')
        
        # Extract key requirements from job description
        key_requirements = []
        if 'react' in job_description:
            key_requirements.append('React')
        if 'angular' in job_description:
            key_requirements.append('Angular')
        if 'node.js' in job_description:
            key_requirements.append('Node.js')
        if 'python' in job_description:
            key_requirements.append('Python')
        if 'mobile' in job_description:
            key_requirements.append('mobile development')
        if 'api' in job_description:
            key_requirements.append('RESTful APIs')
        if 'database' in job_description:
            key_requirements.append('database management')
        
        # Create optimized summary for entry-level
        if 'senior' in job_title.lower() or 'lead' in job_title.lower():
            # For senior positions, be honest about being entry-level but show potential
            optimized_summary = f"Passionate Junior Developer with strong foundation in {', '.join(key_requirements[:3]) if key_requirements else 'modern web technologies'}. "
            optimized_summary += "Seeking first professional opportunity to apply technical skills in building responsive web applications and RESTful APIs. "
            optimized_summary += "Demonstrated ability through personal projects and internship experience. "
            optimized_summary += "Eager to learn, grow, and contribute to a development team while writing clean, maintainable code."
        else:
            # For junior/entry positions, match the title
            optimized_summary = f"Passionate {job_title} with strong foundation in {', '.join(key_requirements[:3]) if key_requirements else 'modern web technologies'}. "
            optimized_summary += "Building responsive web applications, RESTful APIs, and user-friendly interfaces through personal projects and learning. "
            optimized_summary += "Strong problem-solving skills and eagerness to learn in professional environment. "
            optimized_summary += "Committed to writing clean, maintainable code and continuously improving technical skills."
        
        return optimized_summary
    
    def _optimize_skills(self, current_skills: Dict[str, List[str]], job_description: str) -> Dict[str, List[str]]:
        """Optimize skills section for job requirements"""
        optimized_skills = current_skills.copy()
        
        # Add job-specific skills if mentioned in description
        if 'mobile' in job_description:
            optimized_skills['frontend'].extend(['React Native', 'Flutter', 'iOS', 'Android'])
        if 'cloud' in job_description:
            optimized_skills['tools'].extend(['AWS', 'Azure', 'Google Cloud'])
        if 'devops' in job_description:
            optimized_skills['tools'].extend(['Kubernetes', 'Jenkins', 'Docker', 'CI/CD'])
        if 'data' in job_description:
            optimized_skills['other'].extend(['Data Analysis', 'Machine Learning', 'Python'])
        
        # Remove duplicates
        for category in optimized_skills:
            optimized_skills[category] = list(set(optimized_skills[category]))
        
        return optimized_skills
    
    def _optimize_experience(self, current_experience: List[Dict[str, Any]], job_description: str) -> List[Dict[str, Any]]:
        """Optimize experience descriptions for job relevance"""
        optimized_experience = []
        
        for exp in current_experience:
            optimized_exp = exp.copy()
            
            # Add job-relevant keywords to descriptions
            if 'mobile' in job_description and 'mobile' not in exp['description'].lower():
                optimized_exp['description'] += " Experience with mobile application development and responsive design."
            
            if 'api' in job_description and 'api' not in exp['description'].lower():
                optimized_exp['description'] += " Developed and integrated RESTful APIs for seamless data exchange."
            
            if 'database' in job_description and 'database' not in exp['description'].lower():
                optimized_exp['description'] += " Managed database operations and optimized query performance."
            
            optimized_experience.append(optimized_exp)
        
        return optimized_experience
    
    def _optimize_projects(self, current_projects: List[Dict[str, Any]], job_description: str) -> List[Dict[str, Any]]:
        """Optimize project descriptions for job relevance"""
        optimized_projects = []
        
        for project in current_projects:
            optimized_project = project.copy()
            
            # Add job-relevant keywords to project descriptions
            if 'mobile' in job_description and 'mobile' not in project['description'].lower():
                optimized_project['description'] += " Features mobile-responsive design and cross-platform compatibility."
            
            if 'api' in job_description and 'api' not in project['description'].lower():
                optimized_project['description'] += " Integrated with external APIs for enhanced functionality."
            
            if 'database' in job_description and 'database' not in project['description'].lower():
                optimized_project['description'] += " Implemented efficient database design and data management."
            
            optimized_projects.append(optimized_project)
        
        return optimized_projects
    
    def _generate_cv_content(self, cv_data: Dict[str, Any]) -> str:
        """Generate CV content from structured data"""
        content = f"""
{cv_data['header']['name']}
{cv_data['header']['title']}

CONTACT INFORMATION
Email: {cv_data['header']['email']}
Phone: {cv_data['header']['phone']}
Address: {cv_data['header']['address']}
LinkedIn: {cv_data['header']['linkedin']}
GitHub: {cv_data['header']['github']}
Portfolio: {cv_data['header']['portfolio']}

SUMMARY
{cv_data['summary']}

TECHNICAL SKILLS
Frontend Development: {', '.join(cv_data['skills']['frontend'])}
Backend Development: {', '.join(cv_data['skills']['backend'])}
Database & Storage: {', '.join(cv_data['skills']['database'])}
Development Tools: {', '.join(cv_data['skills']['tools'])}
Other Skills: {', '.join(cv_data['skills']['other'])}

WORK EXPERIENCE
"""
        
        for exp in cv_data['experience']:
            content += f"""
{exp['title']} - {exp['company']}
Duration: {exp['duration']}
Description: {exp['description']}
Key Achievements:
"""
            for achievement in exp.get('achievements', []):
                content += f"• {achievement}\n"
        
        content += f"""
EDUCATION
"""
        
        for edu in cv_data['education']:
            content += f"""
{edu['degree']} - {edu['institution']}
Duration: {edu['duration']}
Description: {edu['description']}
"""
        
        content += f"""
PROJECTS
"""
        
        for project in cv_data['projects']:
            content += f"""
{project['name']}
Description: {project['description']}
Technologies: {', '.join(project['technologies'])}
"""
            # Add GitHub code link
            if project.get('code_url'):
                content += f"GitHub Code: {project['code_url']}\n"
            
            # Add live demo link
            if project.get('live_url'):
                content += f"Live Demo: {project['live_url']}\n"
            
            content += "Key Achievements:\n"
            for achievement in project.get('achievements', []):
                content += f"• {achievement}\n"
        
        content += f"""
CERTIFICATES
"""
        
        for cert in cv_data['certificates']:
            content += f"• {cert['name']} - {cert['issuer']} ({cert['date']})\n"
        
        content += f"""
LANGUAGES
"""
        
        for lang in cv_data['languages']:
            content += f"• {lang['language']}: {lang['level']}\n"
        
        return content.strip()
    
    def save_cv_files(self, result: Dict[str, Any], output_dir: str = None) -> Dict[str, str]:
        """Save CV in multiple formats"""
        if not result:
            return {}
        
        if not output_dir:
            timestamp = datetime.now().strftime("%Y%m%d")
            job_title = result['job_info'].get('title', 'developer').replace(' ', '_')
            company = result['job_info'].get('company', 'company').replace(' ', '_')
            output_dir = f"CV_{job_title}_{company}_{timestamp}"
        
        # Create output directory
        Path(output_dir).mkdir(exist_ok=True)
        
        saved_files = {}
        
        # Save text version with professional name
        txt_file = os.path.join(output_dir, "Abdallah_Nasr_Ali_CV.txt")
        with open(txt_file, 'w', encoding='utf-8') as f:
            f.write(result['cv_content'])
        saved_files['txt'] = txt_file
        
        # Save JSON version (for internal use)
        json_file = os.path.join(output_dir, "CV_Data.json")
        with open(json_file, 'w', encoding='utf-8') as f:
            json.dump(result, f, indent=2, ensure_ascii=False)
        saved_files['json'] = json_file
        
        # Save ATS report
        report_file = os.path.join(output_dir, "ATS_Analysis_Report.txt")
        report_content = self.ats_optimizer.generate_ats_report(
            self.ats_optimizer.analyze_cv_for_ats(
                result['cv_content'],
                result['job_info'].get('description', ''),
                result['job_info'].get('title', ''),
                result['job_info'].get('company', '')
            )
        )
        with open(report_file, 'w', encoding='utf-8') as f:
            f.write(report_content)
        saved_files['report'] = report_file
        
        # Try to save PDF version with professional name
        try:
            pdf_file = os.path.join(output_dir, "Abdallah_Nasr_Ali_CV.pdf")
            if self._create_pdf_cv(result['cv_content'], pdf_file):
                saved_files['pdf'] = pdf_file
        except Exception as e:
            self.logger.warning(f"Could not create PDF: {e}")
        
        # Try to save DOCX version with professional name
        try:
            docx_file = os.path.join(output_dir, "Abdallah_Nasr_Ali_CV.docx")
            if self._create_docx_cv(result['cv_content'], docx_file):
                saved_files['docx'] = docx_file
        except Exception as e:
            self.logger.warning(f"Could not create DOCX: {e}")
        
        self.logger.info(f"CV files saved to: {output_dir}")
        return saved_files
    
    def _create_pdf_cv(self, content: str, output_file: str) -> bool:
        """Create PDF version of CV"""
        try:
            from reportlab.lib.pagesizes import letter
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import inch
            
            doc = SimpleDocTemplate(output_file, pagesize=letter)
            styles = getSampleStyleSheet()
            
            # Custom styles
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=16,
                spaceAfter=12,
                alignment=1
            )
            
            heading_style = ParagraphStyle(
                'CustomHeading',
                parent=styles['Heading2'],
                fontSize=14,
                spaceAfter=8,
                spaceBefore=12
            )
            
            normal_style = ParagraphStyle(
                'CustomNormal',
                parent=styles['Normal'],
                fontSize=10,
                spaceAfter=6
            )
            
            # Build content
            story = []
            lines = content.split('\n')
            
            for line in lines:
                line = line.strip()
                if line:
                    if line in ['CONTACT INFORMATION', 'SUMMARY', 'TECHNICAL SKILLS', 'WORK EXPERIENCE', 'EDUCATION', 'PROJECTS', 'CERTIFICATES', 'LANGUAGES']:
                        story.append(Paragraph(line, heading_style))
                        story.append(Spacer(1, 6))
                    elif line.startswith('ABDALLAH NASR ALI'):
                        story.append(Paragraph(line, title_style))
                        story.append(Spacer(1, 12))
                    else:
                        story.append(Paragraph(line, normal_style))
                        story.append(Spacer(1, 4))
            
            doc.build(story)
            return True
            
        except ImportError:
            self.logger.warning("reportlab not available for PDF generation")
            return False
        except Exception as e:
            self.logger.error(f"Error creating PDF: {e}")
            return False
    
    def _create_docx_cv(self, content: str, output_file: str) -> bool:
        """Create DOCX version of CV"""
        try:
            from docx import Document
            from docx.shared import Inches
            
            doc = Document()
            
            # Set margins
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(0.5)
                section.bottom_margin = Inches(0.5)
                section.left_margin = Inches(0.5)
                section.right_margin = Inches(0.5)
            
            # Add content
            lines = content.split('\n')
            for line in lines:
                line = line.strip()
                if line:
                    if line in ['CONTACT INFORMATION', 'SUMMARY', 'TECHNICAL SKILLS', 'WORK EXPERIENCE', 'EDUCATION', 'PROJECTS', 'CERTIFICATES', 'LANGUAGES']:
                        doc.add_heading(line, level=1)
                    elif line.startswith('ABDALLAH NASR ALI'):
                        doc.add_heading(line, level=0)
                    else:
                        doc.add_paragraph(line)
            
            doc.save(output_file)
            return True
            
        except ImportError:
            self.logger.warning("python-docx not available for DOCX generation")
            return False
        except Exception as e:
            self.logger.error(f"Error creating DOCX: {e}")
            return False

def main():
    """Test the ATS CV Generator"""
    print("🎯 ATS CV GENERATOR - Testing Module")
    print("=" * 50)
    
    # Sample job information
    sample_job = {
        "title": "Senior Full Stack Developer",
        "company": "TechCorp",
        "location": "Cairo, Egypt",
        "description": """
        We are looking for a Senior Full Stack Developer to join our team.
        The ideal candidate will have experience with React, Node.js, and modern web technologies.
        Experience with mobile development, APIs, and database management is required.
        Strong problem-solving skills and ability to work in fast-paced environments.
        """
    }
    
    try:
        generator = ATSCVGenerator()
        
        # Generate job-specific CV
        result = generator.generate_job_specific_cv(sample_job)
        
        if result:
            print(f"✅ Job-specific CV generated successfully!")
            print(f"📊 ATS Score: {result['ats_analysis']['overall_score']}/100")
            print(f"🎯 Job: {result['job_info']['title']} at {result['job_info']['company']}")
            
            # Save CV files
            saved_files = generator.save_cv_files(result)
            
            print(f"\n📁 Files created:")
            for format_type, file_path in saved_files.items():
                print(f"   • {format_type.upper()}: {file_path}")
            
            # Display ATS analysis summary
            analysis = result['ats_analysis']
            print(f"\n📊 ATS Analysis Summary:")
            print(f"   • Contact Information: {analysis['contact_info_score']}/100")
            print(f"   • Job Title Match: {analysis['job_title_match_score']}/100")
            print(f"   • Skill Matching: {analysis['skill_match_score']}/100")
            print(f"   • Formatting: {analysis['formatting_score']}/100")
            print(f"   • Readability: {analysis['readability_score']}/100")
            print(f"   • Web Presence: {analysis['web_presence_score']}/100")
            
            if analysis['recommendations']:
                print(f"\n🔧 Top Recommendations:")
                for i, rec in enumerate(analysis['recommendations'][:5], 1):
                    print(f"   {i}. {rec}")
            
        else:
            print("❌ Failed to generate CV")
        
    except Exception as e:
        print(f"❌ Error: {e}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    main()

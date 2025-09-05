#!/usr/bin/env python3
"""
Smart CV Generator - Creates actual different CV versions with real modifications
"""

import os
import json
from datetime import datetime
import PyPDF2
from docx import Document
from docx.shared import Inches

def read_pdf_cv():
    """Read the actual PDF CV"""
    try:
        with open('Abdallah Nasr Ali_Cv.pdf', 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text()
            return text
    except FileNotFoundError:
        print("❌ Abdallah Nasr Ali_Cv.pdf not found!")
        return None
    except Exception as e:
        print(f"❌ Error reading PDF: {e}")
        return None

def create_cv_versions(original_cv):
    """Create 3 actually different versions from the actual CV"""
    
    # Version 1: ATS-Friendly (clean, keyword-rich, keep URLs but clean formatting)
    ats_version = original_cv
    # Clean up extra spaces and formatting for ATS systems
    ats_version = ats_version.replace("  ", " ")
    ats_version = ats_version.replace("  ", " ")  # Double replace to catch all
    # Keep all URLs intact - ATS systems can handle them
    # Just clean up the formatting
    
    # Version 2: Professional (enhanced descriptions, polished language)
    professional_version = original_cv
    professional_version = professional_version.replace("Collaborative and adaptable", "Results-driven and collaborative")
    professional_version = professional_version.replace("passion for clean code", "expertise in clean code architecture")
    professional_version = professional_version.replace("UI/UX", "user interface and user experience design")
    professional_version = professional_version.replace("learning new technologies", "staying current with emerging technologies")
    
    # Version 3: Concise (focused, shorter, streamlined)
    concise_version = original_cv
    # Remove extra spaces and make it more compact
    concise_version = concise_version.replace("  ", " ")
    concise_version = concise_version.replace("  ", " ")  # Double replace
    # Make summary more focused
    concise_version = concise_version.replace("Collaborative and adaptable, with a passion for clean code, UI/UX, and learning new technologies.", "Results-driven developer with expertise in clean code and modern technologies.")
    # Shorten some descriptions
    concise_version = concise_version.replace("Built a responsive Angular e-commerce app with product filtering, cart, checkout, and authentication. Integrated REST APIs and deployed via Vercel.", "Built responsive Angular e-commerce app with product filtering, cart, checkout, and authentication. Integrated REST APIs and deployed via Vercel.")
    
    return {
        "ats": ats_version,
        "professional": professional_version, 
        "concise": concise_version
    }

def create_cv_structure():
    """Create the CV structure with actually different content"""
    
    # Read actual PDF CV
    original_cv = read_pdf_cv()
    if not original_cv:
        return
    
    # Create versions
    cv_versions = create_cv_versions(original_cv)
    
    # Create main folder with timestamp
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    main_folder = f"SMART_CV_{timestamp}"
    
    if not os.path.exists(main_folder):
        os.makedirs(main_folder)
        print(f"📁 Created main folder: {main_folder}")
    
    # Create ALL folder
    all_folder = os.path.join(main_folder, "ALL")
    if not os.path.exists(all_folder):
        os.makedirs(all_folder)
    
    # Create date/time folder
    current_time = datetime.now()
    date_str = current_time.strftime("%Y-%m-%d")
    time_str = current_time.strftime("%I%p")
    
    date_time_folder = os.path.join(all_folder, date_str, time_str)
    if not os.path.exists(date_time_folder):
        os.makedirs(date_time_folder)
    
    # Version folders will be created dynamically for each CV version
    
            # Generate CVs for each version
        for version_name, cv_content in cv_versions.items():
            print(f"\n🔄 Generating {version_name} version...")
            
            # Create version-specific folder
            version_folder = os.path.join(date_time_folder, version_name.upper())
            if not os.path.exists(version_folder):
                os.makedirs(version_folder)
            
            # JSON - clean data only
            json_path = os.path.join(version_folder, f"Abdallah_Nasr_Ali_CV_{timestamp}.json")
            create_clean_json_cv(cv_content, version_name, json_path)
            
            # PDF - clean content only
            pdf_path = os.path.join(version_folder, f"Abdallah_Nasr_Ali_CV_{timestamp}.pdf")
            create_clean_pdf_cv(cv_content, pdf_path)
            
            # DOCX - clean content only
            docx_path = os.path.join(version_folder, f"Abdallah_Nasr_Ali_CV_{timestamp}.docx")
            create_clean_docx_cv(cv_content, docx_path)
            
            # Markdown - clean content only
            md_path = os.path.join(version_folder, f"Abdallah_Nasr_Ali_CV_{timestamp}.md")
            create_clean_markdown_cv(cv_content, md_path)
            
            # Text - clean content only
            txt_path = os.path.join(version_folder, f"Abdallah_Nasr_Ali_CV_{timestamp}.txt")
            create_clean_text_cv(cv_content, txt_path)
    
    print(f"\n🎉 Smart CV versions created successfully!")
    print(f"📁 Main folder: {main_folder}")
    print(f"📁 Structure: ALL/{date_str}/{time_str}/")
    
    # Show differences
    print(f"\n📊 Version differences:")
    for version_name, content in cv_versions.items():
        print(f"  {version_name}: {len(content)} characters")

def create_clean_json_cv(content, version_name, file_path):
    """Create clean JSON CV with structured data"""
    # Split content into sections for better structure
    lines = content.split('\n')
    
    # Extract key information
    cv_data = {
        "version": version_name,
        "personal_info": {
            "name": "Abdallah Nasr Ali",
            "phone": "+201069509757",
            "email": "body16nasr16bn@gmail.com",
            "location": "Cairo, Egypt",
            "github": "https://github.com/AbdalahNasr",
            "portfolio": "https://my-v3-potfolio.vercel.app"
        },
        "summary": "",
        "education": [],
        "skills": {
            "frontend": [],
            "backend": [],
            "tools": []
        },
        "experience": [],
        "projects": [],
        "certificates": [],
        "languages": [],
        "full_content": content
    }
    
    # Parse content to extract structured data
    current_section = ""
    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        if line in ['Summary', 'Experience', 'Education', 'Skills', 'Projects', 'Certificates', 'Languages']:
            current_section = line
        elif current_section == "Summary" and line and not line.startswith('•'):
            cv_data["summary"] += line + " "
        elif current_section == "Education" and line.startswith('•'):
            cv_data["education"].append(line)
        elif current_section == "Skills":
            if "Frontend:" in line:
                skills = line.replace("Frontend:", "").strip().split(", ")
                cv_data["skills"]["frontend"] = [s.strip() for s in skills if s.strip()]
            elif "Backend & Database:" in line:
                skills = line.replace("Backend & Database:", "").strip().split(", ")
                cv_data["skills"]["backend"] = [s.strip() for s in skills if s.strip()]
            elif "Tools:" in line:
                skills = line.replace("Tools:", "").strip().split(", ")
                cv_data["skills"]["tools"] = [s.strip() for s in skills if s.strip()]
        elif current_section == "Experience" and "—" in line:
            cv_data["experience"].append(line)
        elif current_section == "Projects" and "—" in line:
            cv_data["projects"].append(line)
        elif current_section == "Certificates" and line.startswith('•'):
            cv_data["certificates"].append(line)
        elif current_section == "Languages" and line.startswith('•'):
            cv_data["languages"].append(line)
    
    # Clean up summary
    cv_data["summary"] = cv_data["summary"].strip()
    
    with open(file_path, 'w', encoding='utf-8') as f:
        json.dump(cv_data, f, indent=2, ensure_ascii=False)
    print(f"  ✅ JSON: {os.path.basename(file_path)}")

def create_clean_pdf_cv(content, file_path):
    """Create clean PDF CV with proper PDF formatting"""
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        
        # Create PDF document
        doc = SimpleDocTemplate(file_path, pagesize=letter)
        styles = getSampleStyleSheet()
        
        # Create custom styles
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=16,
            spaceAfter=12,
            alignment=1  # Center alignment
        )
        
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=14,
            spaceAfter=8,
            spaceBefore=12
        )
        
        normal_style = ParagraphStyle(
            'CustomNormal',
            parent=styles['Normal'],
            fontSize=10,
            spaceAfter=6
        )
        
        # Build content
        story = []
        
        # Split content into lines
        lines = content.split('\n')
        for line in lines:
            line = line.strip()
            if line:
                # Check for different line types
                if line in ['Summary', 'Experience', 'Education', 'Skills', 'Projects', 'Certificates', 'Languages']:
                    story.append(Paragraph(line, heading_style))
                    story.append(Spacer(1, 6))
                elif line.startswith('Abdallah Nasr Ali'):
                    story.append(Paragraph(line, title_style))
                    story.append(Spacer(1, 12))
                elif line.startswith('+20') or line.startswith('Github:') or line.startswith('Portfolio:'):
                    story.append(Paragraph(line, normal_style))
                    story.append(Spacer(1, 6))
                elif line.startswith('•') or line.startswith('Frontend:') or line.startswith('Backend:') or line.startswith('Tools:'):
                    story.append(Paragraph(line, normal_style))
                    story.append(Spacer(1, 4))
                else:
                    story.append(Paragraph(line, normal_style))
                    story.append(Spacer(1, 4))
        
        # Build PDF
        doc.build(story)
        print(f"  ✅ PDF: {os.path.basename(file_path)}")
        
    except ImportError:
        print(f"  ⚠️  PDF: {os.path.basename(file_path)} (reportlab not available, created text file)")
        # Fallback to text file if reportlab is not available
        with open(file_path, 'w', encoding='utf-8') as f:
            f.write(content)
    except Exception as e:
        print(f"  ❌ PDF: {os.path.basename(file_path)} - Error: {e}")
        # Fallback to text file on error
        with open(file_path, 'w', encoding='utf-8') as f:
            f.write(content)

def create_clean_docx_cv(content, file_path):
    """Create clean DOCX CV with proper formatting"""
    doc = Document()
    
    # Set margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
    
    # Split content into lines and add to document
    lines = content.split('\n')
    for line in lines:
        line = line.strip()
        if line:  # Only add non-empty lines
            # Check if it's a header (name, contact info, section titles)
            if line in ['Summary', 'Experience', 'Education', 'Skills', 'Projects']:
                doc.add_heading(line, level=1)
            elif line.startswith('Abdallah Nasr Ali') or line.startswith('+20') or line.startswith('Github:') or line.startswith('Portfolio:'):
                doc.add_paragraph(line)
            else:
                doc.add_paragraph(line)
    
    # Save the document
    doc.save(file_path)
    print(f"  ✅ DOCX: {os.path.basename(file_path)}")

def create_clean_markdown_cv(content, file_path):
    """Create clean Markdown CV with just the content"""
    with open(file_path, 'w', encoding='utf-8') as f:
        f.write(content)
    print(f"  ✅ Markdown: {os.path.basename(file_path)}")

def create_clean_text_cv(content, file_path):
    """Create clean Text CV with just the content"""
    with open(file_path, 'w', encoding='utf-8') as f:
        f.write(content)
    print(f"  ✅ Text: {os.path.basename(file_path)}")

if __name__ == "__main__":
    create_cv_structure()

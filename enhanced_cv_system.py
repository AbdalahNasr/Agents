#!/usr/bin/env python3
"""
Enhanced CV System - Multi-Stack ATS-Friendly CV Generator

This system creates professional, ATS-friendly CVs for different tech stacks:
- Frontend Developer
- Backend Developer  
- Full-Stack Developer
- DevOps Engineer
- Data Scientist
- Mobile Developer
- UI/UX Designer

Features:
- ATS-optimized formatting
- Stack-specific content customization
- Multiple export formats (TXT, JSON, MD, PDF, DOCX)
- Keyword optimization for job matching
- Professional templates
"""

import json
import os
from datetime import datetime
from typing import Dict, List, Optional
import openai

from config import Config

class EnhancedCVSystem:
    """Advanced CV system for multiple tech stacks with ATS optimization."""
    
    def __init__(self):
        """Initialize the enhanced CV system."""
        self.openai_available = bool(Config.OPENAI_API_KEY)
        if self.openai_available:
            openai.api_key = Config.OPENAI_API_KEY
        
        # Tech stack definitions with keywords and focus areas
        self.tech_stacks = {
            'frontend': {
                'name': 'Frontend Developer',
                'keywords': ['React', 'Vue', 'Angular', 'JavaScript', 'TypeScript', 'HTML', 'CSS', 'SASS', 'Webpack', 'Responsive Design'],
                'focus_areas': ['User Experience', 'Performance', 'Accessibility', 'Cross-browser Compatibility'],
                'common_roles': ['Frontend Developer', 'UI Developer', 'JavaScript Developer', 'React Developer']
            },
            'backend': {
                'name': 'Backend Developer',
                'keywords': ['Python', 'Java', 'C#', 'Node.js', 'Go', 'Rust', 'SQL', 'NoSQL', 'APIs', 'Microservices'],
                'focus_areas': ['System Architecture', 'Database Design', 'API Development', 'Performance Optimization'],
                'common_roles': ['Backend Developer', 'API Developer', 'Software Engineer', 'Systems Developer']
            },
            'fullstack': {
                'name': 'Full-Stack Developer',
                'keywords': ['React', 'Node.js', 'Python', 'JavaScript', 'TypeScript', 'MongoDB', 'PostgreSQL', 'AWS', 'Docker'],
                'focus_areas': ['End-to-End Development', 'System Integration', 'Full Application Lifecycle'],
                'common_roles': ['Full-Stack Developer', 'Software Engineer', 'Web Developer', 'Application Developer']
            },
            'devops': {
                'name': 'DevOps Engineer',
                'keywords': ['Docker', 'Kubernetes', 'AWS', 'Azure', 'GCP', 'Jenkins', 'GitLab CI', 'Terraform', 'Ansible'],
                'focus_areas': ['Infrastructure as Code', 'CI/CD Pipelines', 'Cloud Architecture', 'Monitoring & Logging'],
                'common_roles': ['DevOps Engineer', 'Site Reliability Engineer', 'Infrastructure Engineer', 'Cloud Engineer']
            },
            'data_science': {
                'name': 'Data Scientist',
                'keywords': ['Python', 'R', 'SQL', 'Pandas', 'NumPy', 'Scikit-learn', 'TensorFlow', 'PyTorch', 'Jupyter'],
                'focus_areas': ['Machine Learning', 'Statistical Analysis', 'Data Visualization', 'Predictive Modeling'],
                'common_roles': ['Data Scientist', 'ML Engineer', 'Data Analyst', 'Research Scientist']
            },
            'mobile': {
                'name': 'Mobile Developer',
                'keywords': ['React Native', 'Flutter', 'iOS', 'Android', 'Swift', 'Kotlin', 'Mobile UI/UX', 'App Store'],
                'focus_areas': ['Mobile Performance', 'Cross-platform Development', 'Native Features', 'App Store Optimization'],
                'common_roles': ['Mobile Developer', 'iOS Developer', 'Android Developer', 'App Developer']
            },
            'ui_ux': {
                'name': 'UI/UX Designer',
                'keywords': ['Figma', 'Sketch', 'Adobe XD', 'User Research', 'Wireframing', 'Prototyping', 'Design Systems'],
                'focus_areas': ['User Experience Design', 'Visual Design', 'Interaction Design', 'User Research'],
                'common_roles': ['UI/UX Designer', 'Product Designer', 'Interaction Designer', 'Visual Designer']
            }
        }
        
        # ATS optimization rules
        self.ats_rules = {
            'formatting': [
                'Use standard fonts (Arial, Calibri, Times New Roman)',
                'Keep font size between 10-12pt',
                'Use clear section headers',
                'Avoid tables and complex formatting',
                'Use bullet points for lists',
                'Keep consistent spacing'
            ],
            'keywords': [
                'Include industry-standard job titles',
                'Use specific technology names',
                'Include quantifiable achievements',
                'Use action verbs',
                'Avoid abbreviations unless common'
            ],
            'content': [
                'Keep sections concise and focused',
                'Use numbers and metrics',
                'Highlight relevant experience',
                'Customize for each application',
                'Proofread for errors'
            ]
        }
    
    def generate_stack_specific_cv(self, base_cv: str, target_stack: str, 
                                 job_title: str = None, company: str = None) -> Dict:
        """
        Generate a CV optimized for a specific tech stack.
        
        Args:
            base_cv: Original CV text
            target_stack: Target tech stack (frontend, backend, fullstack, etc.)
            job_title: Specific job title if known
            company: Company name if known
            
        Returns:
            Dict: Stack-specific CV with multiple versions
        """
        if target_stack not in self.tech_stacks:
            raise ValueError(f"Unknown tech stack: {target_stack}. Available: {list(self.tech_stacks.keys())}")
        
        stack_info = self.tech_stacks[target_stack]
        
        print(f"🎯 Generating CV for {stack_info['name']} role...")
        
        # Parse base CV
        cv_sections = self._parse_cv_sections(base_cv)
        
        # Customize for target stack
        customized_sections = self._customize_for_stack(cv_sections, stack_info, job_title, company)
        
        # Create ATS-optimized version
        ats_version = self._create_ats_version(customized_sections, stack_info)
        
        # Create professional version
        professional_version = self._create_professional_version(customized_sections, stack_info)
        
        # Create concise version
        concise_version = self._create_concise_version(customized_sections, stack_info)
        
        # Generate stack-specific suggestions
        suggestions = self._generate_stack_suggestions(customized_sections, stack_info)
        
        result = {
            'stack': target_stack,
            'stack_info': stack_info,
            'original_cv': base_cv,
            'customized_sections': customized_sections,
            'versions': {
                'ats_optimized': ats_version,
                'professional': professional_version,
                'concise': concise_version
            },
            'suggestions': suggestions,
            'generated_at': datetime.now().isoformat(),
            'job_title': job_title,
            'company': company
        }
        
        return result
    
    def _parse_cv_sections(self, cv_text: str) -> Dict[str, str]:
        """Parse CV text into logical sections."""
        sections = {}
        
        # Simple section parsing
        section_patterns = {
            'summary': r'(?:summary|objective|profile|about)\s*:?\s*\n(.*?)(?=\n\s*\n|\n\s*[A-Z]|\Z)',
            'experience': r'(?:experience|work\s+history|employment)\s*:?\s*\n(.*?)(?=\n\s*\n|\n\s*[A-Z]|\Z)',
            'skills': r'(?:skills|technical\s+skills|competencies)\s*:?\s*\n(.*?)(?=\n\s*\n|\n\s*[A-Z]|\Z)',
            'education': r'(?:education|academic|qualifications)\s*:?\s*\n(.*?)(?=\n\s*\n|\n\s*[A-Z]|\Z)',
            'projects': r'(?:projects|portfolio|achievements)\s*:?\s*\n(.*?)(?=\n\s*\n|\n\s*[A-Z]|\Z)'
        }
        
        for section_name, pattern in section_patterns.items():
            import re
            match = re.search(pattern, cv_text, re.IGNORECASE | re.DOTALL)
            if match:
                sections[section_name] = match.group(1).strip()
            else:
                sections[section_name] = self._extract_section_by_content(cv_text, section_name)
        
        return sections
    
    def _extract_section_by_content(self, cv_text: str, section_name: str) -> str:
        """Extract section content based on keywords and context."""
        if section_name == 'summary':
            lines = cv_text.split('\n')
            for line in lines[:5]:
                if line.strip() and len(line.strip()) > 50:
                    return line.strip()
        elif section_name == 'experience':
            experience_keywords = ['worked', 'employed', 'job', 'position', 'role', 'responsibilities']
            lines = cv_text.split('\n')
            experience_lines = []
            for line in lines:
                if any(keyword in line.lower() for keyword in experience_keywords):
                    experience_lines.append(line.strip())
            return '\n'.join(experience_lines[:10])
        elif section_name == 'skills':
            skill_keywords = ['python', 'javascript', 'java', 'sql', 'aws', 'docker', 'react', 'node']
            lines = cv_text.split('\n')
            skill_lines = []
            for line in lines:
                if any(skill in line.lower() for skill in skill_keywords):
                    skill_lines.append(line.strip())
            return '\n'.join(skill_lines)
        return ""
    
    def _customize_for_stack(self, cv_sections: Dict[str, str], stack_info: Dict, 
                            job_title: str = None, company: str = None) -> Dict[str, str]:
        """Customize CV sections for specific tech stack."""
        customized = {}
        
        for section_name, content in cv_sections.items():
            if section_name == 'summary':
                customized[section_name] = self._customize_summary(content, stack_info, job_title, company)
            elif section_name == 'skills':
                customized[section_name] = self._customize_skills(content, stack_info)
            elif section_name == 'experience':
                customized[section_name] = self._customize_experience(content, stack_info)
            elif section_name == 'projects':
                customized[section_name] = self._customize_projects(content, stack_info)
            else:
                customized[section_name] = content
        
        return customized
    
    def _customize_summary(self, summary: str, stack_info: Dict, job_title: str = None, company: str = None) -> str:
        """Customize summary for specific tech stack."""
        role = job_title or stack_info['name']
        
        # Create stack-specific summary
        stack_summary = f"{role} with expertise in {', '.join(stack_info['keywords'][:5])}. "
        stack_summary += f"Specialized in {', '.join(stack_info['focus_areas'][:2])}. "
        
        if summary:
            # Extract key achievements from original summary
            if 'experience' in summary.lower():
                stack_summary += "Proven track record of delivering high-quality software solutions. "
            if 'team' in summary.lower() or 'lead' in summary.lower():
                stack_summary += "Experienced in leading development teams and mentoring junior developers. "
        
        stack_summary += f"Passionate about creating innovative solutions using modern {stack_info['name'].lower()} technologies."
        
        return stack_summary
    
    def _customize_skills(self, skills: str, stack_info: Dict) -> str:
        """Customize skills section for specific tech stack."""
        if not skills:
            return ', '.join(stack_info['keywords'])
        
        # Reorganize skills to prioritize stack-specific ones
        stack_keywords = [kw.lower() for kw in stack_info['keywords']]
        original_skills = skills.split(',')
        
        # Separate stack-specific and general skills
        stack_skills = []
        general_skills = []
        
        for skill in original_skills:
            skill_lower = skill.strip().lower()
            if any(keyword.lower() in skill_lower for keyword in stack_keywords):
                stack_skills.append(skill.strip())
            else:
                general_skills.append(skill.strip())
        
        # Create organized skills section
        organized_skills = []
        if stack_skills:
            organized_skills.append(f"Core Technologies: {', '.join(stack_skills[:8])}")
        if general_skills:
            organized_skills.append(f"Additional Skills: {', '.join(general_skills[:6])}")
        
        return ' | '.join(organized_skills)
    
    def _customize_experience(self, experience: str, stack_info: Dict) -> str:
        """Customize experience section for specific tech stack."""
        if not experience:
            return experience
        
        # Highlight stack-relevant experience
        lines = experience.split('\n')
        customized_lines = []
        
        for line in lines:
            if line.strip():
                # Highlight stack-specific keywords
                for keyword in stack_info['keywords']:
                    if keyword.lower() in line.lower():
                        line = line.replace(keyword, f"**{keyword}**")
                        break
                customized_lines.append(line)
        
        return '\n'.join(customized_lines)
    
    def _customize_projects(self, projects: str, stack_info: Dict) -> str:
        """Customize projects section for specific tech stack."""
        if not projects:
            return projects
        
        # Highlight stack-relevant projects
        lines = projects.split('\n')
        customized_lines = []
        
        for line in lines:
            if line.strip():
                # Highlight stack-specific technologies
                for keyword in stack_info['keywords']:
                    if keyword.lower() in line.lower():
                        line = line.replace(keyword, f"**{keyword}**")
                        break
                customized_lines.append(line)
        
        return '\n'.join(customized_lines)
    
    def _create_ats_version(self, sections: Dict[str, str], stack_info: Dict) -> str:
        """Create ATS-optimized version of CV."""
        ats_cv = []
        
        # Header
        ats_cv.append("PROFESSIONAL SUMMARY")
        ats_cv.append("=" * 50)
        ats_cv.append(sections.get('summary', ''))
        ats_cv.append("")
        
        # Skills (most important for ATS)
        ats_cv.append("TECHNICAL SKILLS")
        ats_cv.append("=" * 50)
        ats_cv.append(sections.get('skills', ''))
        ats_cv.append("")
        
        # Experience
        ats_cv.append("PROFESSIONAL EXPERIENCE")
        ats_cv.append("=" * 50)
        ats_cv.append(sections.get('experience', ''))
        ats_cv.append("")
        
        # Projects
        ats_cv.append("KEY PROJECTS")
        ats_cv.append("=" * 50)
        ats_cv.append(sections.get('projects', ''))
        ats_cv.append("")
        
        # Education
        ats_cv.append("EDUCATION")
        ats_cv.append("=" * 50)
        ats_cv.append(sections.get('education', ''))
        
        return '\n'.join(ats_cv)
    
    def _create_professional_version(self, sections: Dict[str, str], stack_info: Dict) -> str:
        """Create professional version of CV."""
        prof_cv = []
        
        # Professional header
        prof_cv.append(f"{stack_info['name'].upper()}")
        prof_cv.append("=" * len(stack_info['name']))
        prof_cv.append("")
        
        # Summary
        prof_cv.append("PROFESSIONAL SUMMARY")
        prof_cv.append("-" * 20)
        prof_cv.append(sections.get('summary', ''))
        prof_cv.append("")
        
        # Core competencies
        prof_cv.append("CORE COMPETENCIES")
        prof_cv.append("-" * 20)
        prof_cv.append(f"• {', '.join(stack_info['focus_areas'])}")
        prof_cv.append("")
        
        # Skills
        prof_cv.append("TECHNICAL EXPERTISE")
        prof_cv.append("-" * 20)
        prof_cv.append(sections.get('skills', ''))
        prof_cv.append("")
        
        # Experience
        prof_cv.append("PROFESSIONAL EXPERIENCE")
        prof_cv.append("-" * 20)
        prof_cv.append(sections.get('experience', ''))
        prof_cv.append("")
        
        # Projects
        prof_cv.append("NOTABLE PROJECTS")
        prof_cv.append("-" * 20)
        prof_cv.append(sections.get('projects', ''))
        prof_cv.append("")
        
        # Education
        prof_cv.append("EDUCATION & CERTIFICATIONS")
        prof_cv.append("-" * 20)
        prof_cv.append(sections.get('education', ''))
        
        return '\n'.join(prof_cv)
    
    def _create_concise_version(self, sections: Dict[str, str], stack_info: Dict) -> str:
        """Create concise version of CV."""
        concise_cv = []
        
        # Brief header
        concise_cv.append(f"{stack_info['name']}")
        concise_cv.append("")
        
        # Brief summary
        summary = sections.get('summary', '')
        if len(summary) > 200:
            summary = summary[:200] + "..."
        concise_cv.append(summary)
        concise_cv.append("")
        
        # Key skills only
        concise_cv.append("Key Skills: " + sections.get('skills', '')[:100] + "...")
        concise_cv.append("")
        
        # Experience highlights
        experience = sections.get('experience', '')
        if len(experience) > 300:
            experience = experience[:300] + "..."
        concise_cv.append("Experience: " + experience)
        
        return '\n'.join(concise_cv)
    
    def _generate_stack_suggestions(self, sections: Dict[str, str], stack_info: Dict) -> List[str]:
        """Generate stack-specific improvement suggestions."""
        suggestions = []
        
        # Stack-specific suggestions
        suggestions.append(f"Highlight {stack_info['name']} experience prominently")
        suggestions.append(f"Emphasize {', '.join(stack_info['focus_areas'][:2])} skills")
        suggestions.append(f"Include {', '.join(stack_info['keywords'][:5])} in experience descriptions")
        
        # General suggestions
        suggestions.append("Use quantifiable achievements (e.g., 'increased performance by 40%')")
        suggestions.append("Include specific project outcomes and metrics")
        suggestions.append("Tailor content to each job application")
        suggestions.append("Keep ATS-friendly formatting")
        
        return suggestions
    
    def export_cv(self, cv_data: Dict, format_type: str = 'text', filename_prefix: str = None) -> str:
        """
        Export CV to different formats.
        
        Args:
            cv_data: CV data from generate_stack_specific_cv
            format_type: Export format (text, json, markdown, pdf, docx)
            filename_prefix: Custom filename prefix
            
        Returns:
            str: File path or content
        """
        if not filename_prefix:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename_prefix = f"cv_{cv_data['stack']}_{timestamp}"
        
        if format_type == 'json':
            return self._export_to_json(cv_data, filename_prefix)
        elif format_type == 'markdown':
            return self._export_to_markdown(cv_data, filename_prefix)
        elif format_type == 'pdf':
            return self._export_to_pdf(cv_data, filename_prefix)
        elif format_type == 'docx':
            return self._export_to_docx(cv_data, filename_prefix)
        else:
            return self._export_to_text(cv_data, filename_prefix)
    
    def _export_to_json(self, cv_data: Dict, filename_prefix: str) -> str:
        """Export CV to JSON format."""
        filename = f"{filename_prefix}.json"
        with open(filename, 'w', encoding='utf-8') as f:
            json.dump(cv_data, f, indent=2, ensure_ascii=False)
        return filename
    
    def _export_to_markdown(self, cv_data: Dict, filename_prefix: str) -> str:
        """Export CV to Markdown format."""
        filename = f"{filename_prefix}.md"
        
        md_content = f"# {cv_data['stack_info']['name']} CV\n\n"
        md_content += f"**Generated:** {cv_data['generated_at']}\n"
        md_content += f"**Target Stack:** {cv_data['stack']}\n\n"
        
        # Add versions
        for version_name, content in cv_data['versions'].items():
            md_content += f"## {version_name.replace('_', ' ').title()}\n\n"
            md_content += f"```\n{content}\n```\n\n"
        
        # Add suggestions
        md_content += "## Improvement Suggestions\n\n"
        for suggestion in cv_data['suggestions']:
            md_content += f"- {suggestion}\n"
        
        with open(filename, 'w', encoding='utf-8') as f:
            f.write(md_content)
        
        return filename
    
    def _export_to_text(self, cv_data: Dict, filename_prefix: str) -> str:
        """Export CV to text format."""
        filename = f"{filename_prefix}.txt"
        
        text_content = f"{cv_data['stack_info']['name']} CV\n"
        text_content += "=" * (len(cv_data['stack_info']['name']) + 6) + "\n\n"
        
        # Add versions
        for version_name, content in cv_data['versions'].items():
            text_content += f"{version_name.upper()}\n"
            text_content += "-" * len(version_name) + "\n"
            text_content += f"{content}\n\n"
        
        # Add suggestions
        text_content += "IMPROVEMENT SUGGESTIONS\n"
        text_content += "-" * 25 + "\n"
        for suggestion in cv_data['suggestions']:
            text_content += f"• {suggestion}\n"
        
        with open(filename, 'w', encoding='utf-8') as f:
            f.write(text_content)
        
        return filename
    
    def _export_to_pdf(self, cv_data: Dict, filename_prefix: str) -> str:
        """Export CV to PDF format."""
        try:
            from reportlab.lib.pagesizes import letter
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            
            filename = f"{filename_prefix}.pdf"
            
            doc = SimpleDocTemplate(filename, pagesize=letter)
            story = []
            styles = getSampleStyleSheet()
            
            # Title
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=18,
                spaceAfter=30,
                alignment=1
            )
            story.append(Paragraph(f"{cv_data['stack_info']['name']} CV", title_style))
            story.append(Spacer(1, 20))
            
            # Add ATS version content
            ats_content = cv_data['versions']['ats_optimized']
            lines = ats_content.split('\n')
            
            for line in lines:
                if line.startswith('='):
                    # Section header
                    section_name = line.replace('=', '').strip()
                    if section_name:
                        story.append(Paragraph(section_name, styles['Heading2']))
                        story.append(Spacer(1, 12))
                elif line.strip():
                    story.append(Paragraph(line.strip(), styles['Normal']))
                    story.append(Spacer(1, 6))
            
            doc.build(story)
            return filename
            
        except ImportError:
            return "PDF export failed - reportlab not installed"
        except Exception as e:
            return f"PDF export failed: {str(e)}"
    
    def _export_to_docx(self, cv_data: Dict, filename_prefix: str) -> str:
        """Export CV to DOCX format."""
        try:
            from docx import Document
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            
            filename = f"{filename_prefix}.docx"
            
            doc = Document()
            
            # Title
            title = doc.add_heading(f"{cv_data['stack_info']['name']} CV", 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add ATS version content
            ats_content = cv_data['versions']['ats_optimized']
            lines = ats_content.split('\n')
            
            for line in lines:
                if line.startswith('='):
                    # Section header
                    section_name = line.replace('=', '').strip()
                    if section_name:
                        doc.add_heading(section_name, level=1)
                elif line.strip():
                    doc.add_paragraph(line.strip())
            
            doc.save(filename)
            return filename
            
        except ImportError:
            return "DOCX export failed - python-docx not installed"
        except Exception as e:
            return f"DOCX export failed: {str(e)}"


def main():
    """Main function to demonstrate the enhanced CV system."""
    print("🚀 Enhanced CV System - Multi-Stack ATS-Friendly CV Generator")
    print("=" * 70)
    
    # Read the sample CV
    try:
        with open('sample_cv.txt', 'r', encoding='utf-8') as f:
            base_cv = f.read()
        print(f"✅ Loaded base CV ({len(base_cv)} characters)")
    except FileNotFoundError:
        print("❌ sample_cv.txt not found. Please create it first.")
        return
    
    # Initialize CV system
    cv_system = EnhancedCVSystem()
    
    # Available tech stacks
    print(f"\n🎯 Available Tech Stacks:")
    for stack, info in cv_system.tech_stacks.items():
        print(f"  • {stack}: {info['name']}")
    
    # Generate CV for different stacks
    target_stacks = ['frontend', 'backend', 'fullstack', 'devops']
    
    for stack in target_stacks:
        print(f"\n{'='*50}")
        print(f"Generating CV for {stack.upper()} stack...")
        
        try:
            # Generate stack-specific CV
            cv_result = cv_system.generate_stack_specific_cv(
                base_cv, stack, 
                job_title=f"{cv_system.tech_stacks[stack]['name']}",
                company="Tech Company"
            )
            
            print(f"✅ Generated CV for {stack}")
            print(f"📊 Versions: {len(cv_result['versions'])}")
            print(f"💡 Suggestions: {len(cv_result['suggestions'])}")
            
            # Export to different formats
            print(f"\n📁 Exporting CV...")
            
            # Text export
            text_file = cv_system.export_cv(cv_result, 'text', f"cv_{stack}")
            print(f"  ✅ Text: {text_file}")
            
            # JSON export
            json_file = cv_system.export_cv(cv_result, 'json', f"cv_{stack}")
            print(f"  ✅ JSON: {json_file}")
            
            # Markdown export
            md_file = cv_system.export_cv(cv_result, 'markdown', f"cv_{stack}")
            print(f"  ✅ Markdown: {md_file}")
            
            # PDF export
            pdf_file = cv_system.export_cv(cv_result, 'pdf', f"cv_{stack}")
            print(f"  ✅ PDF: {pdf_file}")
            
            # DOCX export
            docx_file = cv_system.export_cv(cv_result, 'docx', f"cv_{stack}")
            print(f"  ✅ DOCX: {docx_file}")
            
        except Exception as e:
            print(f"❌ Failed to generate CV for {stack}: {e}")
    
    print(f"\n🎉 CV generation completed!")
    print(f"\n📋 Generated files:")
    for file in os.listdir('.'):
        if file.startswith('cv_') and file.endswith(('.txt', '.json', '.md', '.pdf', '.docx')):
            print(f"  📄 {file}")


if __name__ == "__main__":
    main()
